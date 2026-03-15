# Shared test fixtures: temp config, sample files, mock embedder.

from pathlib import Path
from unittest.mock import MagicMock, patch

import numpy as np
import pytest

from smart_search.config import SmartSearchConfig


@pytest.fixture
def tmp_config(tmp_path):
    """SmartSearchConfig with all paths pointing to tmp_path."""
    return SmartSearchConfig(
        lancedb_path=str(tmp_path / "vectors"),
        sqlite_path=str(tmp_path / "metadata.db"),
    )


@pytest.fixture
def sample_pdf_path(tmp_path):
    """Generate a minimal 2-page PDF with reportlab.

    Page 1: heading + body paragraph.
    Page 2: a simple table.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    pdf_path = tmp_path / "test_document.pdf"
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    styles = getSampleStyleSheet()

    elements = [
        Paragraph("Chapter 1: Introduction to Machine Learning", styles["Heading1"]),
        Spacer(1, 12),
        Paragraph(
            "Machine learning is a subset of artificial intelligence that enables "
            "systems to learn and improve from experience without being explicitly "
            "programmed. This document covers the fundamentals of supervised and "
            "unsupervised learning approaches.",
            styles["Normal"],
        ),
        Spacer(1, 400),  # Force page break
        Paragraph("Chapter 2: Results Summary", styles["Heading1"]),
        Spacer(1, 12),
        Table(
            [
                ["Model", "Accuracy", "F1 Score"],
                ["Random Forest", "0.92", "0.91"],
                ["SVM", "0.89", "0.87"],
                ["Neural Network", "0.95", "0.94"],
            ],
            style=TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]),
        ),
    ]

    doc.build(elements)
    return pdf_path


@pytest.fixture
def sample_docx_path(tmp_path):
    """Generate a minimal DOCX with python-docx.

    Contains a heading and one body paragraph.
    """
    from docx import Document

    docx_path = tmp_path / "test_document.docx"
    doc = Document()
    doc.add_heading("Regulatory Compliance Overview", level=1)
    doc.add_paragraph(
        "This document outlines the regulatory requirements for medical device "
        "software under FDA 21 CFR Part 820 and EU MDR 2017/745. All manufacturers "
        "must establish and maintain a quality management system."
    )
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def mock_embedder(tmp_config):
    """Mock Embedder producing deterministic vectors matching config dims."""
    mock = MagicMock()
    rng = np.random.RandomState(42)
    dims = tmp_config.embedding_dimensions

    def fake_embed_documents(texts):
        return [rng.randn(dims).tolist() for _ in texts]

    def fake_embed_query(query):
        return rng.randn(dims).tolist()

    mock.embed_documents.side_effect = fake_embed_documents
    mock.embed_query.side_effect = fake_embed_query
    mock.get_model_name.return_value = tmp_config.embedding_model
    mock.is_loaded = True
    return mock
